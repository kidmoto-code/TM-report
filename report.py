import os
import docx
from docx.shared import Pt, Mm
from loguru import logger
from typing import List, Optional, Tuple
from data import TradeMark, Application


class UserData:
    """
    Класс для хранения данных пользователя с сохранением порядка ввода.

    :ivar ordered_data: Список кортежей (тип, значение) в порядке добавления
    :type ordered_data: List[Tuple[str, str]]
    """

    ordered_data: List[Tuple[str, str]] = []

    @classmethod
    @logger.catch
    def clear(cls) -> None:
        """Очищает все данные пользователя."""
        logger.debug("CLEAR USERDATA")
        cls.ordered_data = []
        cls.clean_image_dir()

    @classmethod
    @logger.catch
    def clean_image_dir(cls, directory="images"):
        """
        Очищает указанную директорию от всех файлов.
        По умолчанию директория "images"

        :param directory str: Путь к директории для очистки
        """
        try:
            # Проверяем, существует ли дирректория
            if not os.path.exists(directory):
                logger.error(f"Директория {directory} не существует")
                return

            # Перебираем все элементы в дирректории
            for filename in os.listdir(directory):
                file_path = os.path.join(directory, filename)

                try:
                    # Если это файл или символическая ссылка - удаляем
                    if os.path.isfile(file_path) or os.path.islink(file_path):
                        os.unlink(file_path)

                except Exception as e:
                    logger.error(f"Не удалось удалить {file_path}.\n{e}")

            logger.info(f"Директория {directory} успешно очищена.")

        except Exception as e:
            logger.error(f"Ошибка при очистке дирректории {directory}.\n{e}")

    @classmethod
    @logger.catch
    def add_data(cls, data: List[str]) -> str:
        """
        Добавляет данные, автоматически определяя их тип.

        :param data List[str]: Список строк с данными для добавления
        :return str: Строка с отчетом о добавленных данных
        """
        logger.debug("ADDING INPUTED DATA")

        for item in data:
            if len(item) == 2 and item.isdigit() and int(item) <= 45:
                cls.ordered_data.append(('MKTU', item))
            elif len(item) == 10 and item.isdigit() and item[:2] == '20':
                cls.ordered_data.append(('APP', item))
            elif 3 <= len(item) <= 9 and item[:-2].isdigit():
                cls.ordered_data.append(('TM', item))
            elif item[0] == "M":
                processed = ("0" * (8 - len(item)) + item[1:]) if len(item) < 7 else item[1:]
                cls.ordered_data.append(('MADRID', processed))
            else:
                cls.ordered_data.append(('ERR', item))

        return cls.answer()

    @classmethod
    def get_data_by_type(cls, data_type: str) -> List[str]:
        """
        Возвращает данные указанного типа.

        :param data_type str: Тип данных ('MKTU', 'TM', 'APP', 'MADRID', 'ERR')
        :return List[str]: Список значений указанного типа
        :raises ValueError: Если передан недопустимый тип данных
        """
        valid_types = {'MKTU', 'TM', 'APP', 'MADRID', 'ERR'}
        if data_type not in valid_types:
            raise ValueError(f"Неверный тип. Допустимые: {valid_types}")

        return [value for (type_, value) in cls.ordered_data if type_ == data_type]

    @classmethod
    def answer(cls) -> str:
        """
        Формирует текстовый отчет о добавленных данных.

        :return str: Форматированная строка с количеством данных каждого типа
        """
        counts = {
            'MKTU': 0,
            'TM': 0,
            'APP': 0,
            'MADRID': 0,
            'ERR': 0
        }

        for type_, _ in cls.ordered_data:
            counts[type_] += 1

        report = [
            f"Ошибок: {counts['ERR']}",
            f"МКТУ: {counts['MKTU']}",
            f"Товарные знаки (TM): {counts['TM']}",
            f"Заявки (APP): {counts['APP']}",
            f"Международные знаки (MADRID): {counts['MADRID']}"
        ]

        if counts['ERR'] > 0:
            report.append(f"Ошибочные данные: {cls.get_data_by_type('ERR')}")

        return "\n".join(report)


class Report:
    """
    Класс для генерации отчета в формате DOCX.

    :ivar doc: Объект документа Word
    :vartype doc: docx.Document
    :ivar style: Стиль оформления документа
    :vartype style: docx.styles.Style
    :ivar tm_data: Список номеров товарных знаков
    :vartype tm_data: List[str]
    :ivar app_data: Список номеров заявок
    :vartype app_data: List[str]
    :ivar madrid_data: Список международных регистраций
    :vartype madrid_data: List[str]
    :ivar mktu_data: Список классов МКТУ
    :vartype mktu_data: List[str]
    """

    def __init__(self) -> None:
        """Инициализирует объект отчета и запускает его формирование."""
        self.doc = docx.Document()
        self.style = self.doc.styles['Normal']
        self.style.font.name = 'Arial'
        self.style.font.size = Pt(11)
        self.style.paragraph_format.first_line_indent = Mm(10)

        self.tm_data = UserData.get_data_by_type('TM')
        self.app_data = UserData.get_data_by_type('APP')
        self.madrid_data = UserData.get_data_by_type('MADRID')
        self.mktu_data = UserData.get_data_by_type('MKTU')
        self.ordered_data = UserData.ordered_data
        self.order = "by_type"  # by_user or by_type

        self.write_docx()
        self.doc.save('report.docx')

    def write_docx(self):
        if self.order == "by_type":
            for idx, tm_number in enumerate(self.tm_data, 1):
                logger.info(f'Обработка ТЗ №{tm_number}')
                self.write_TM(tm_number, idx)

            for idx, app_number in enumerate(self.app_data,
                                             len(self.tm_data)):
                logger.info(f'Обработка Заявки №{app_number}')
                self.write_APP(app_number, idx)

    @logger.catch
    def write_TM(self, tm_number: str, idx: int = 1) -> None:
        """
        Формирует раздел отчета с товарными знаками.

        Для каждого товарного знака добавляет:
            - Изображение знака
            - Основные реквизиты (номер, даты)
            - Правообладателя
            - Классы МКТУ
            - Дополнительную информацию (если имеется)

        :raises Exception: Логирует ошибки обработки отдельных ТЗ
        """

        try:
            tm = TradeMark(tm_number, 'TM')

            if not tm.is_available:
                UserData.ordered_data.append(('ERR', tm_number))
                raise ValueError(f"Документ {tm_number} не найден в ФИПС")

            if not tm.image_path or not os.path.exists(tm.image_path):
                logger.warning(f'Изображение для ТЗ {tm_number} не найдено')
                raise ValueError(f"Изображение для ТЗ {tm_number} не найдено")

            par = self.doc.add_paragraph()
            run = par.add_run()
            run.add_text(f'{idx}. Товарный знак ')
            run.add_picture(tm.image_path)
            run.add_text(f' свидетельство № {tm.number}, ')
            run.add_text(f'дата подачи: {tm.applicationdate}, ')
            run.add_text(f'дата регистрации: {tm.registrationdate}, ')
            run.add_text(f'{tm.holdername}, ')
            run.add_text(f'классы МКТУ: {", ".join(tm.classes[1])}.\n')

            for class_full, class_num in zip(*tm.classes):
                if class_num in self.mktu_data:
                    self.doc.add_paragraph().add_run(f'{class_full}.')

            if tm.unprotected:
                p = self.doc.add_paragraph()
                p.add_run('Примечание! ').bold = True
                p.add_run(f'Неохраняемые элементы: {tm.unprotected}')

            if tm.representative and 'городисский' in tm.representative.lower():
                p = self.doc.add_paragraph()
                p.add_run('ВНИМАНИЕ! ').bold = True
                p.add_run(tm.representative)

        except Exception as e:
            logger.error(f'Ошибка при обработке ТЗ {tm_number}: {str(e)}')
            UserData.ordered_data.append(('ERR', tm_number))

    @logger.catch
    def write_APP(self, app_number: str, idx: int = 1):
        try:
            app = Application(app_number, 'TMAP')

            if not app.is_available:
                UserData.ordered_data.append(('ERR', app_number))
                raise ValueError(f"Документ {app_number} не найден в ФИПС")

            if not app.image_path or not os.path.exists(app.image_path):
                logger.warning(f'Изображение для Заявки {app_number} не найдено')
                raise ValueError(f"Изображение для Заявки {app_number} не найдено")

            par = self.doc.add_paragraph()
            run = par.add_run()
            run.add_text(f'{idx}. Заявленное обозначение ')
            run.add_picture(app.image_path)
            run.add_text(f' заявка № {app.number}, ')
            run.add_text(f'дата подачи заявки: {app.applicationdate}, ')
            run.add_text(f'{app.holdername}, ')
            run.add_text(f'классы МКТУ: {", ".join(app.classes[1])}.\n')

            for class_full, class_num in zip(*app.classes):
                if class_num in self.mktu_data:
                    self.doc.add_paragraph().add_run(f'{class_full}.')

            if app.unprotected:
                p = self.doc.add_paragraph()
                p.add_run('Примечание! ').bold = True
                p.add_run(f'Неохраняемые элементы: {app.unprotected}')

            if app.representative and 'городисский' in app.representative.lower():
                p = self.doc.add_paragraph()
                p.add_run('ВНИМАНИЕ! ').bold = True
                p.add_run(app.representative)

        except Exception as e:
            logger.error(f'Ошибка при обработке ТЗ {app_number}: {str(e)}')
            UserData.ordered_data.append(('ERR', app_number))
